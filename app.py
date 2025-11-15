"""
Career Fair Companion ‚Äî COMPLETE REWRITE (Crash‚Äësafe)
Robust research ‚Ä¢ %-match job search (ATS + company domains) ‚Ä¢ Optional AI (safely gated)
Sidebar tabs (Inputs / Settings / Bulk) ‚Ä¢ Theme presets ‚Ä¢ Resume tailoring ‚Ä¢ Contacts ‚Ä¢ QR card

This version is **crash‚Äësafe when `streamlit` isn‚Äôt installed**: it falls back to a tiny shim so
the module can import and our unit tests can run. In a real deployment you should still
install `streamlit` (add it to `requirements.txt`) to get the full UI.
"""
from __future__ import annotations
import os, io, re, json, difflib
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

import pandas as pd

# -------------------- Streamlit import (with shim fallback) --------------------
try:  # normal path
    import streamlit as st  # type: ignore
    _STREAMLIT_AVAILABLE = True
except ModuleNotFoundError:  # crash-safe path for sandbox/test runners
    _STREAMLIT_AVAILABLE = False

    class _ShimContainer:
        def __enter__(self):
            return self
        def __exit__(self, exc_type, exc, tb):
            return False
        def __getattr__(self, name):
            def _noop(*args, **kwargs):
                # return sensible defaults for common widgets
                if name in ("text_input", "text_area", "selectbox", "color_picker", "radio"):
                    return ""
                if name in ("button", "toggle", "checkbox"):
                    return False
                return None
            return _noop

    class _ShimST:
        def __init__(self):
            self.session_state: Dict[str, Any] = {}
            self.secrets: Dict[str, Any] = {}
            self.sidebar = _ShimContainer()
        # Explicit shims for APIs that must exist
        def spinner(self, *args, **kwargs):
            return _ShimContainer()
        def container(self, *args, **kwargs):
            return _ShimContainer()
        def expander(self, *args, **kwargs):
            return _ShimContainer()
        def tabs(self, labels):
            # Accept list/tuple or int
            try:
                if isinstance(labels, int):
                    n = max(1, int(labels))
                else:
                    n = max(1, int(len(list(labels))))
            except Exception:
                n = 1
            return [_ShimContainer() for _ in range(n)]
        def columns(self, spec=None, *args, **kwargs):
            # Streamlit accepts an int (n columns) or a sequence of ratios
            if spec is None:
                n = 2
            elif isinstance(spec, int):
                n = max(1, spec)
            else:
                try:
                    n = max(1, len(list(spec)))
                except Exception:
                    n = 2
            return [_ShimContainer() for _ in range(n)]
        # Generic no-ops
        def __getattr__(self, name):
            def _noop(*args, **kwargs):
                # No-op UI calls
                if name in (
                    "file_uploader", "link_button", "image", "download_button",
                    "dataframe", "markdown", "write", "caption", "info", "warning",
                    "success", "error", "divider", "set_page_config", "subheader",
                ):
                    return None
                # Inputs with safe defaults
                if name in ("text_input", "text_area", "selectbox", "color_picker", "radio"):
                    return ""
                if name in ("button", "toggle", "checkbox"):
                    return False
                return None
            return _noop

    st = _ShimST()  # type: ignore
    os.environ["NON_UI_MODE"] = "1"  # used by tests below

# -------------------- Optional deps (fail-soft) --------------------
try:
    from duckduckgo_search import DDGS
except Exception:
    DDGS = None
try:
    import wikipedia
except Exception:
    wikipedia = None
try:
    import yake
except Exception:
    yake = None
try:
    import segno
except Exception:
    segno = None
try:
    from openai import OpenAI
except Exception:
    OpenAI = None
try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None
try:
    from docx import Document
except Exception:
    Document = None
try:
    from bs4 import BeautifulSoup  # type: ignore
except Exception:
    BeautifulSoup = None  # type: ignore
import requests

# -------------------- Paths & storage --------------------
DATA_DIR = "data"
RESUME_DIR = os.path.join(DATA_DIR, "resumes")
BULK_FILE = os.path.join(DATA_DIR, "companies.txt")
INDEX_CSV = os.path.join(RESUME_DIR, "index.csv")
os.makedirs(RESUME_DIR, exist_ok=True)

# discover preset lists in /data
PRESET_TXT_FILES: List[str] = []
try:
    if os.path.isdir(DATA_DIR):
        PRESET_TXT_FILES = [os.path.join(DATA_DIR, f) for f in sorted(os.listdir(DATA_DIR)) if f.endswith(".txt")]
except Exception:
    PRESET_TXT_FILES = []

# -------------------- Branding & themes --------------------
DEFAULT_PRIMARY = "#7C3AED"  # violet-600
DEFAULT_ACCENT  = "#22D3EE"  # cyan-400
DEFAULT_CARD_BG = "#0f172a"  # slate-900

SCHOOL_THEMES = {
    "Cal State LA": {"primary": "#000000", "accent": "#FFC72C"},
    "UCLA":         {"primary": "#2774AE", "accent": "#FFD100"},
    "USC":          {"primary": "#990000", "accent": "#FFCC00"},
    "UC Berkeley":  {"primary": "#003262", "accent": "#FDB515"},
    "MIT":          {"primary": "#A31F34", "accent": "#8A8B8C"},
}

# Prefer the corporate entity when ambiguous (e.g., "Jacobs" ‚Üí "Jacobs Solutions")
COMPANY_FIXUPS = {
    "jacobs": "Jacobs Solutions",
    "jacobs engineering": "Jacobs Solutions",
    "pwc": "PwC",
    "ernst & young": "EY",
}

def canonical_company(name: str) -> str:
    key = re.sub(r"[^a-z0-9]+", " ", (name or "").lower()).strip()
    return COMPANY_FIXUPS.get(key, name)

# Known ATS hosts
ATS_HOSTS = [
    "lever.co","greenhouse.io","myworkdayjobs.com","smartrecruiters.com","icims.com",
    "taleo.net","successfactors.com","ashbyhq.com","workable.com","jobvite.com",
    "eightfold.ai","recruitee.com","workforcenow.adp.com","adp.com","breezy.hr",
    "oraclecloud.com","dayforcehcm.com","ultipro.com","bamboohr.com","jazzhr.com",
]

# -------------------- Page & global styles --------------------
st.set_page_config(page_title="Career Fair Companion", page_icon="üéì", layout="wide", initial_sidebar_state="expanded")

st.markdown(
    f"""
    <style>
      @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;800&display=swap');
      html, body, [class*="css"] {{ font-family: Inter, system-ui, -apple-system, Segoe UI, Roboto, sans-serif; }}
      .hero {{ background: linear-gradient(135deg, {DEFAULT_PRIMARY}22, {DEFAULT_ACCENT}22); padding: 1rem 1.25rem; border-radius: 18px; border: 1px solid #1f2937; }}
      .app-title {{ background: linear-gradient(90deg, {DEFAULT_PRIMARY}, {DEFAULT_ACCENT}); -webkit-background-clip: text; background-clip: text; color: transparent; font-weight: 800; letter-spacing:-0.3px; }}
      .card {{ border-radius: 16px; padding: 1rem 1.25rem; border: 1px solid #1f2937; background: {DEFAULT_CARD_BG}; box-shadow: 0 8px 24px rgba(0,0,0,0.25); }}
      .muted {{ color: #94a3b8; }} .small {{ font-size: 0.92rem; }}
      .match-pill {{ display:inline-block; padding: 2px 10px; border-radius: 999px; background: #0ea5e933; border:1px solid #22d3ee; font-weight:600; }}
      .btn-primary button {{ background: {DEFAULT_PRIMARY} !important; border-color: {DEFAULT_PRIMARY} !important; }}
      .btn-accent button {{ background: {DEFAULT_ACCENT} !important; border-color: {DEFAULT_ACCENT} !important; color:#0b1020!important; }}
    </style>
    """,
    unsafe_allow_html=True,
)

def apply_theme(primary: str, accent: str):
    st.markdown(
        f"""
        <style>
          .app-title {{ background: linear-gradient(90deg, {primary}, {accent}); -webkit-background-clip: text; background-clip: text; color: transparent; }}
          .btn-primary button {{ background: {primary} !important; border-color: {primary} !important; }}
          .btn-accent button {{ background: {accent} !important; border-color: {accent} !important; color:#0b1020!important; }}
          .match-pill {{ background: {accent}22; border-color: {accent}; }}
        </style>
        """,
        unsafe_allow_html=True,
    )

# -------------------- AI helpers (safely gated) --------------------

def get_ai_client() -> Optional["OpenAI"]:
    if OpenAI is None:
        return None
    api_key = st.secrets.get("openai_api_key") or os.environ.get("OPENAI_API_KEY")
    base_url = st.secrets.get("openai_base_url") or os.environ.get("OPENAI_BASE_URL")
    if not api_key:
        return None
    try:
        if base_url:
            return OpenAI(api_key=api_key, base_url=base_url)
        return OpenAI(api_key=api_key)
    except Exception:
        return None

def ai_enabled() -> bool:
    if not st.session_state.get("use_ai_toggle"):
        return False
    return bool(st.secrets.get("openai_api_key") or os.environ.get("OPENAI_API_KEY") or st.secrets.get("openai_base_url") or os.environ.get("OPENAI_BASE_URL"))

def summarize_ai(text: str) -> Optional[str]:
    try:
        client = get_ai_client()
        if not client:
            return None
        mdl = st.secrets.get("openai_model") or os.environ.get("OPENAI_MODEL") or "gpt-4o-mini"
        out = client.chat.completions.create(
            model=mdl,
            temperature=0.4,
            messages=[
                {"role":"system","content":"You write crisp, factual company snapshots for students preparing for career fairs."},
                {"role":"user","content":f"Summarize in 4-5 sentences:\n\n{text}"}
            ],
        )
        return (out.choices[0].message.content or "").strip()
    except Exception:
        return None

# -------------------- Search utilities --------------------

def ddg_search(query: str, max_results: int = 10, timelimit: Optional[str] = None) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    if DDGS is None:
        return rows
    try:
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=max_results, timelimit=timelimit, safesearch="moderate"):
                rows.append({"title": r.get("title"), "href": r.get("href"), "body": r.get("body")})
    except Exception:
        pass
    return rows

def clean_snippet(text: Optional[str], length: int = 250) -> str:
    if not text:
        return ""
    t = re.sub(r"\s+", " ", str(text)).strip()
    return (t if len(t) <= length else t[: length - 1].rsplit(" ", 1)[0] + "‚Ä¶")

def fetch_text(url: str, limit: int = 6000) -> str:
    try:
        headers = {"User-Agent": "Mozilla/5.0 (CareerFairComp/1.0)"}
        r = requests.get(url, headers=headers, timeout=8)
        if r.status_code != 200:
            return ""
        html = r.text
        if BeautifulSoup is not None:
            soup = BeautifulSoup(html, "html.parser")
            txt = soup.get_text(" ", strip=True)
        else:
            txt = re.sub(r"<[^>]+>", " ", html)
            txt = re.sub(r"\s+", " ", txt).strip()
        return clean_snippet(txt, length=limit)
    except Exception:
        return ""

def wiki_summary(company: str) -> Optional[str]:
    if not wikipedia:
        return None
    try:
        wikipedia.set_lang("en")
        query = canonical_company(company)
        candidates = [f"{query} (company)", f"{query} (corporation)", query]
        page_title = None
        for q in candidates:
            res = wikipedia.search(q, results=1)
            if res:
                page_title = res[0]
                break
        if not page_title:
            return None
        page = wikipedia.page(page_title, auto_suggest=False)
        # Avoid wrong entities (e.g., biblical Jacob)
        if any(w in page.title.lower() for w in ["patriarch", "given name", "biblical"]):
            return None
        return clean_snippet(page.summary, 600)
    except Exception:
        return None

def find_official_site(company: str) -> Optional[str]:
    for r in ddg_search(f"{company} official site", max_results=4, timelimit="y"):
        url = r.get("href") or ""
        if not url:
            continue
        host = re.sub(r"^https?://", "", url).split("/")[0]
        if not any(s in host for s in ["wikipedia.org", "linkedin.com", "twitter.com"]):
            return url
    return None

COMMON_CAREERS_PATHS = ["/careers", "/jobs", "/join-us", "/careers/search", "/en/careers"]

def find_careers_candidates(company: str) -> List[str]:
    cand = []
    site = find_official_site(company)
    if site:
        base = re.sub(r"/$", "", site)
        cand += [base + p for p in COMMON_CAREERS_PATHS]
    for r in ddg_search(f"{company} careers", max_results=6, timelimit="y"):
        u = r.get("href");
        if u and u not in cand:
            cand.append(u)
    return cand[:10]

def categorize_domain(url: str) -> str:
    try:
        host = re.sub(r"^https?://", "", url).split("/")[0]
    except Exception:
        return "Other"
    if any(k in host for k in ATS_HOSTS):
        return "ATS"
    if any(k in host for k in ["indeed.com", "linkedin.com/jobs", "glassdoor.com/job"]):
        return "Job Board"
    if any(k in host for k in ["newsroom", "press", "media", "blog."]):
        return "Press/Blog"
    return "Other"

# -------------------- Source preference + de-duplication --------------------

def get_source_bonus(src: str) -> float:
    """Map source -> bonus weight based on user preference.
    Value is multiplied by 10 in the final score.
    Profiles:
      - ATS-first: ATS=1.0, Job Board=0.6, Other=0.4 (default)
      - Balanced:  ATS=0.8, Job Board=0.8, Other=0.4
      - Discovery: ATS=0.8, Job Board=1.0, Other=0.4
    """
    pref = st.session_state.get("source_pref", "ATS-first")
    if pref == "Balanced":
        weights = {"ATS": 0.8, "Job Board": 0.8, "Other": 0.4}
    elif pref == "Discovery":
        weights = {"ATS": 0.8, "Job Board": 1.0, "Other": 0.4}
    else:  # ATS-first
        weights = {"ATS": 1.0, "Job Board": 0.6, "Other": 0.4}
    return weights.get(src, 0.4)

# De-duplication helper: prefer ATS when multiple entries share the same title

def _norm_title_key(title: str) -> str:
    t = (title or "").lower()
    t = re.sub(r"\b(internship|intern|new grad|graduate|early career)\b", "", t)
    t = re.sub(r"[^a-z0-9]+", " ", t).strip()
    return t


def dedupe_jobs(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """De-duplicate by URL and normalized title; prefer ATS over Job Board over Other."""
    seen_urls: set = set()
    best_by_title: Dict[str, Dict[str, Any]] = {}
    def rank(src: str) -> int:
        return {"ATS": 2, "Job Board": 1, "Other": 0}.get(src, 0)
    for r in rows:
        u = r.get("url")
        if not u or u in seen_urls:
            continue
        seen_urls.add(u)
        key = _norm_title_key(r.get("title", ""))
        prev = best_by_title.get(key)
        if prev is None or rank(r.get("source", "Other")) > rank(prev.get("source", "Other")):
            best_by_title[key] = r
    return list(best_by_title.values())

# -------------------- Core job aggregator --------------------

def job_search(company: str, role: str = "", location: str = "", max_results: int = 30) -> List[Dict[str, Any]]:
    brand = canonical_company(company)
    seeds = [
        "site:lever.co","site:greenhouse.io","site:myworkdayjobs.com","site:smartrecruiters.com","site:icims.com",
        "site:taleo.net","site:successfactors.com","site:ashbyhq.com","site:workable.com","site:jobvite.com",
        "site:eightfold.ai","site:recruitee.com","site:workforcenow.adp.com","site:adp.com","site:breezy.hr",
        "site:oraclecloud.com","site:dayforcehcm.com","site:ultipro.com","site:bamboohr.com","site:jazzhr.com",
    ]
    terms = f"{brand} {role}".strip()
    if not role:
        terms += " intern OR internship OR graduate"
    if location:
        terms += f" {location}"

    rows: List[Dict[str, Any]] = []
    # ATS sweep
    for s in seeds:
        q = f"{s} {terms}"
        for r in ddg_search(q, max_results=6, timelimit="m"):
            rows.append({"title": r["title"], "url": r["href"], "snippet": clean_snippet(r["body"]), "source": categorize_domain(r["href"])})
    # Company domain sweep
    for cand in find_careers_candidates(brand):
        try:
            host = re.sub(r"^https?://", "", cand).split("/")[0]
            q = f"site:{host} {role or 'intern OR internship OR graduate'}"
            if location:
                q += f" {location}"
            for r in ddg_search(q, max_results=6, timelimit="m"):
                rows.append({"title": r["title"], "url": r["href"], "snippet": clean_snippet(r["body"]), "source": categorize_domain(r["href"])})
        except Exception:
            continue
    # Last‚Äëditch heuristic
    if not rows:
        for r in ddg_search(f"{brand} {role or 'internship'} apply", max_results=8, timelimit="m"):
            rows.append({"title": r["title"], "url": r["href"], "snippet": clean_snippet(r["body"]), "source": categorize_domain(r["href"])})

    # De‚Äëdup & limit
    uniq = dedupe_jobs(rows)
    return uniq[:max_results]

# -------------------- %-match scoring --------------------

def _norm_tokens(s: str) -> List[str]:
    s = s.lower(); s = re.sub(r"[^a-z0-9\+\-\s]", " ", s)
    return [t for t in s.split() if len(t) > 2]

def score_job(job: Dict[str, Any], role: str, interests: List[str], location: str) -> Tuple[int, str]:
    title = job.get("title") or ""; snip = job.get("snippet") or ""; url = job.get("url") or ""
    body = snip or fetch_text(url)
    role_ratio = difflib.SequenceMatcher(a=(role or '').lower(), b=(title + " " + body).lower()).ratio() if role else 0.0
    body_tokens = set(_norm_tokens(title + " " + body))
    interests_norm = [i.lower().strip() for i in interests if i.strip()]
    overlap = sum(1 for i in interests_norm if any(tok in body_tokens for tok in _norm_tokens(i)))
    interest_ratio = (overlap / max(1, len(interests_norm))) if interests_norm else 0.0
    loc_hit = 1 if (location and location.lower() in (title + " " + body + " " + url).lower()) else 0
    src = job.get("source", "Other"); source_bonus = get_source_bonus(src)
    score = 40 * role_ratio + 40 * interest_ratio + 10 * loc_hit + 10 * source_bonus
    score = max(0, min(100, int(round(score))))
    reason = f"role~{int(role_ratio*100)}%, interests {overlap}/{len(interests_norm) or 1}, location {'‚úì' if loc_hit else '‚Äî'}, source {src}"
    return score, reason

# -------------------- Keyword extractor --------------------

def extract_keywords(text: str, top_k: int = 12) -> List[str]:
    t = (text or "").strip()
    if not t: return []
    if yake:
        try:
            kw = yake.KeywordExtractor(top=top_k)
            return [k for k,_ in kw.extract_keywords(t)][:top_k]
        except Exception:
            pass
    toks = [w for w in re.findall(r"[A-Za-z][A-Za-z0-9\-\+]{2,}", t.lower()) if w not in {"and","the","with","for","from","that","this","will","your","you"}]
    freq: Dict[str, int] = {}
    for x in toks: freq[x] = freq.get(x, 0) + 1
    return [k for k,_ in sorted(freq.items(), key=lambda kv: kv[1], reverse=True)][:top_k]

# -------------------- Resume helpers --------------------

def parse_resume(upload) -> Tuple[str, bytes, str]:
    raw = upload.read(); name = upload.name.lower(); ext = ".txt"
    if name.endswith(".pdf"):
        ext = ".pdf"
        if PdfReader is None: return "", raw, ext
        try:
            reader = PdfReader(io.BytesIO(raw))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n".join(pages), raw, ext
        except Exception: return "", raw, ext
    if name.endswith(".docx"):
        ext = ".docx"
        if Document is None: return "", raw, ext
        try:
            doc = Document(io.BytesIO(raw))
            txt = "\n".join([p.text for p in doc.paragraphs])
            return txt, raw, ext
        except Exception: return "", raw, ext
    try:
        return raw.decode("utf-8", errors="ignore"), raw, ".txt"
    except Exception:
        return "", raw, ".bin"

def save_docx_from_text(text: str, path: str):
    if Document is None: return False
    try:
        doc = Document()
        for line in text.splitlines():
            if line.strip().startswith("# "): doc.add_heading(line.strip("# ").strip(), level=1)
            elif line.strip().startswith("## "): doc.add_heading(line.strip("# ").strip(), level=2)
            else: doc.add_paragraph(line)
        doc.save(path); return True
    except Exception: return False

def append_resume_index(meta: Dict[str, Any]):
    cols = ["timestamp","company","role","job_url","file"]
    row = {k: meta.get(k, "") for k in cols}
    df = pd.DataFrame([row])
    if os.path.exists(INDEX_CSV):
        base = pd.read_csv(INDEX_CSV)
        base = pd.concat([base, df], ignore_index=True)
        base.to_csv(INDEX_CSV, index=False)
    else:
        df.to_csv(INDEX_CSV, index=False)

# -------------------- Contacts & QR --------------------

def find_contacts(company: str, role: str = "") -> List[Dict[str, str]]:
    if not company: return []
    queries = [
        f'site:linkedin.com/in ("Recruiter" OR "Talent Acquisition" OR "University Recruiter") "{company}"',
        f'site:linkedin.com/company {company} people recruiter',
        f'{company} university recruiting email',
        f'{company} HR contact careers',
    ]
    rows: List[Dict[str, str]] = []
    for q in queries:
        for r in ddg_search(q, max_results=6, timelimit="y"):
            rows.append({"title": r.get("title") or "", "url": r.get("href") or "", "snippet": clean_snippet(r.get("body")), "source": categorize_domain(r.get("href") or "")})
    seen = set(); uniq = []
    for row in rows:
        u = row["url"]
        if u in seen: continue
        seen.add(u); uniq.append(row)
    return uniq

def build_vcard(name: str, title: str, email: str, phone: str, org: str, linkedin: str, website: str) -> str:
    lines = ["BEGIN:VCARD","VERSION:3.0",f"FN:{name}", f"TITLE:{title}" if title else "", f"ORG:{org}" if org else "", f"EMAIL;TYPE=INTERNET:{email}" if email else "", f"TEL;TYPE=CELL:{phone}" if phone else "", f"URL:{linkedin}" if linkedin else "", f"URL:{website}" if website else "", "END:VCARD"]
    return "\n".join([l for l in lines if l])

# -------------------- Bulk companies helpers --------------------

def extract_domain(url: str) -> str:
    try:
        url = url.strip()
        if not re.match(r"^https?://", url):
            url = "https://" + url
        host = re.sub(r"^https?://", "", url).split("/")[0]
        return host.lower()
    except Exception:
        return url

def ddg_find_careers(domain: str) -> List[Dict[str, str]]:
    queries = [f"site:{domain} careers", f"site:{domain} jobs", f"site:{domain} internships", f"site:{domain} early careers"]
    hits = []
    for q in queries:
        for r in ddg_search(q, max_results=5, timelimit="y"):
            hits.append({"title": r.get("title"), "url": r.get("href"), "snippet": clean_snippet(r.get("body"))})
    seen = set(); uniq = []
    for h in hits:
        u = h["url"]
        if u in seen: continue
        seen.add(u); uniq.append(h)
    return uniq

# -------------------- Session defaults --------------------
for key, default in {
    "projects": [], "jobs": [], "questions": {}, "wiki": None,
    "saved": [], "contacts": [],
    "resume_text": "", "resume_bytes": b"", "resume_ext": "",
    "assistant_history": [],
    "bulk_urls": [], "bulk_results": [],
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# -------------------- Sidebar --------------------
with st.sidebar:
    tabs = st.tabs(["Inputs", "Settings", "Bulk"])  # Bulk now in sidebar

    with tabs[0]:
        st.subheader("üéì Major Preset + Research")
        MAJOR_PRESETS = {
            "Civil Engineering": {"role": "Civil Engineering Intern","interests": ["transportation","water","sustainability","Python"]},
            "Computer Science / Software": {"role": "Software Engineer Intern","interests": ["backend","web","AI","data"]},
            "Data / Analytics": {"role": "Data Analyst Intern","interests": ["analytics","dashboards","SQL"]},
            "Nursing": {"role": "Student Nurse / Nursing Intern","interests": ["clinical","patient care","EMR"]},
            "(Custom)": {"role": "Intern","interests": ["projects","leadership"]},
        }
        major = st.selectbox("Major (preset)", options=list(MAJOR_PRESETS.keys()), index=0)
        if st.button("Apply Preset"):
            p = MAJOR_PRESETS.get(major, MAJOR_PRESETS["(Custom)"])
            st.session_state.role_input = p.get("role", "Intern")
            st.session_state.interests_input = ", ".join(p.get("interests", []))
        company = st.text_input("Company name", placeholder="e.g., Jacobs, AECOM, Google")
        role = st.text_input("Target role/title", value=st.session_state.get("role_input","Intern"))
        location = st.text_input("Location filter (optional)", placeholder="e.g., Los Angeles, CA")
        interests_str = st.text_input("Interests / keywords (comma-separated)", value=st.session_state.get("interests_input",""))
        interests = [s.strip() for s in interests_str.split(",") if s.strip()]
        colA, colB = st.columns(2)
        with colA: go_research = st.button("Research", type="primary")
        with colB: go_questions = st.button("Make Questions")
        colC, colD = st.columns(2)
        with colC: go_jobs = st.button("Find Jobs")
        with colD: go_export = st.button("Export Sheet")

    with tabs[1]:
        st.subheader("‚öôÔ∏è Settings")
        school = st.selectbox("School theme", options=["(Default)"] + list(SCHOOL_THEMES.keys()), index=0)
        if school != "(Default)":
            t = SCHOOL_THEMES.get(school, {})
            primary = t.get("primary", DEFAULT_PRIMARY); accent = t.get("accent", DEFAULT_ACCENT)
        else:
            primary = DEFAULT_PRIMARY; accent = DEFAULT_ACCENT
        custom_primary = st.color_picker("Primary", value=primary)
        custom_accent  = st.color_picker("Accent", value=accent)
        apply_theme(custom_primary or primary, custom_accent or accent)
        st.markdown("---")
        st.subheader("üîé Search ranking")
        _pref_options = ["ATS-first", "Balanced", "Discovery"]
        _pref_default = st.session_state.get("source_pref", "ATS-first")
        source_pref = st.selectbox(
            "Source preference",
            options=_pref_options,
            index=_pref_options.index(_pref_default) if _pref_default in _pref_options else 0,
            help="Control ranking bias: ATS-first (precise + direct apply), Balanced, or Discovery (broader via job boards).",
        )
        st.session_state["source_pref"] = source_pref
        st.markdown("---")
        st.subheader("ü§ñ AI Settings")
        use_ai = st.toggle("Use AI (OpenAI or compatible)", value=False)
        st.session_state["use_ai_toggle"] = use_ai
        st.text_input("Model (optional)", value=os.environ.get("OPENAI_MODEL", ""))
        st.text_input("Base URL (optional)", value=os.environ.get("OPENAI_BASE_URL", ""))

    with tabs[2]:
        st.subheader("üì¶ Bulk Companies (quick)")
        txt_sb = st.text_area("Paste URLs (one per line)", height=110)
        preset_labels = [f"data/{os.path.basename(p)}" for p in PRESET_TXT_FILES] or ["(none)"]
        preset_choice = st.selectbox("Load preset list", preset_labels, index=0)
        if st.button("Load preset in sidebar") and PRESET_TXT_FILES:
            chosen_fp = PRESET_TXT_FILES[preset_labels.index(preset_choice)]
            try:
                with open(chosen_fp, "r", encoding="utf-8") as f:
                    st.session_state.bulk_urls = [l.strip() for l in f if l.strip() and not l.strip().startswith('#')]
                st.success(f"Loaded {len(st.session_state.bulk_urls)} URLs from {os.path.basename(chosen_fp)}")
            except Exception as e:
                st.error(f"Could not load preset: {e}")
        if st.button("Scan (sidebar)"):
            urls = []
            if txt_sb.strip():
                urls += [u.strip() for u in txt_sb.splitlines() if u.strip() and not u.strip().startswith('#')]
            if st.session_state.get("bulk_urls"):
                urls += [u for u in st.session_state.bulk_urls if u]
            urls = list(dict.fromkeys(urls))[:60]
            if not urls:
                st.warning("No URLs provided. Paste or load a preset.")
            else:
                items = []
                for u in urls:
                    dom = extract_domain(u)
                    hits = ddg_find_careers(dom)
                    if not hits:
                        items.append({"domain": dom, "careers_url": "(none)", "title": "", "snippet": ""})
                    else:
                        for h in hits[:2]:
                            items.append({"domain": dom, "careers_url": h["url"], "title": h["title"], "snippet": h["snippet"]})
                st.session_state.bulk_results = items
                st.success("Scanned. See ‚ÄòBulk Companies‚Äô tab for full table/CSV.")

# -------------------- Header --------------------
st.markdown(
    """
    <div class="hero">
      <div style="display:flex; align-items:center; gap:12px;">
        <div style="font-size:2.2rem;">üéì</div>
        <div>
          <div class="app-title" style="font-size:2rem;">Career Fair Companion</div>
          <div class="muted small">Research companies fast. Ask sharper questions. Find roles. Tailor resumes. Swap QR cards.</div>
        </div>
      </div>
    </div>
    """,
    unsafe_allow_html=True,
)

# -------------------- Research / Questions / Jobs triggers --------------------
if DDGS is None:
    st.warning("Search library (duckduckgo-search) is unavailable. Ensure it's in requirements.txt and installed.")

if 'go_research' not in globals():
    go_research = go_questions = go_jobs = go_export = False

if go_research and 'company' in locals() and company:
    with st.spinner("Researching‚Ä¶"):
        cn = canonical_company(company)
        snap = wiki_summary(cn)
        if not snap:
            site = find_official_site(cn)
            if site:
                snap = fetch_text(site, limit=500)
        st.session_state.wiki = snap or "(No summary found yet. Try Recent Projects or Jobs.)"
        if ai_enabled() and st.session_state.wiki and "No summary" not in st.session_state.wiki:
            s = summarize_ai(st.session_state.wiki)
            if s: st.session_state.wiki = s
    st.success("Company research updated.")

if go_questions:
    st.session_state.questions = {
        "Role & Impact": [
            f"How does {locals().get('role') or 'this role'} define success in the first 90 days?",
            "Which projects are top priority this quarter?",
            "What metrics matter most for early‚Äëcareer hires?",
        ],
        "Team & Mentorship": [
            "What does onboarding look like?",
            "How often are 1:1s and feedback cycles?",
            "Any examples of cross‚Äëteam collaboration for interns/new grads?",
        ],
        "Projects & Methods": [
            "Which tools are standard here?",
            "How do you scope projects on tight timelines?",
            f"For someone into {', '.join((locals().get('interests') or [])[:5]) or 'the core skills for this role'}, what would you recommend studying?",
        ],
        "Culture & Growth": [
            "How do interns find mentors and stretch projects?",
            "Are there rotations or formal training?",
            "How do you support learning (certs, conferences)?",
        ],
        "Recruiter‚ÄëSpecific": [
            "Which teams should I watch based on my interests?",
            "Common pitfalls you see in applications?",
            "What should I include in a follow‚Äëup email after the fair?",
        ],
    }
    st.success("Questions ready.")

if go_jobs and 'company' in locals() and company:
    with st.spinner("Searching for jobs‚Ä¶"):
        base_jobs = job_search(company, role=locals().get('role', ''), location=locals().get('location', ''))
        scored = []
        for j in base_jobs:
            pct, why = score_job(j, locals().get('role', ''), locals().get('interests', []), locals().get('location', ''))
            j2 = {**j, "match": pct, "why": why}
            scored.append(j2)
        st.session_state.jobs = sorted(scored, key=lambda x: x.get("match", 0), reverse=True)
    st.success("Job results updated.")

# -------------------- Main tabs --------------------
overview_tab, jobs_tab, bulk_tab, apply_tab, contacts_tab, qr_tab, saved_tab, help_tab = st.tabs([
    "Overview","Jobs","Bulk Companies","Apply","Contacts","My Card (QR)","Saved","Help & Setup",
])

with overview_tab:
    if 'company' not in locals() or not company:
        st.info("Enter a company in the sidebar and click **Research**.")
    else:
        st.subheader(company)
        if st.session_state.wiki:
            st.markdown("**Company Snapshot**")
            st.write(st.session_state.wiki)
        else:
            st.caption("No snapshot yet.")

with jobs_tab:
    st.subheader("Job Listings (ATS + official domains)")
    c1, c2, c3 = st.columns([1.2,1,1])
    company_in = c1.text_input("Company", value=locals().get('company',''), placeholder="e.g., Jacobs, AECOM, Google")
    role_in    = c2.text_input("Role keywords", value=locals().get('role',''), placeholder="e.g., Civil Engineering Intern")
    loc_in     = c3.text_input("Location (optional)", value=locals().get('location',''), placeholder="City, State")
    if st.button("Search Jobs", type="primary"):
        with st.spinner("Searching across ATS and the company's domains‚Ä¶"):
            rows = job_search(company_in, role=role_in, location=loc_in)
            scored = []
            interests_here = locals().get('interests', [])
            for j in rows:
                pct, why = score_job(j, role_in, interests_here, loc_in)
                scored.append({**j, "match": pct, "why": why})
            st.session_state.jobs = sorted(scored, key=lambda x: x.get("match", 0), reverse=True)
    st.divider()
    if st.session_state.jobs:
        sort_match = st.checkbox("Sort by % match", value=True)
        jobs_view = sorted(st.session_state.jobs, key=lambda x: x.get("match", 0), reverse=True) if sort_match else st.session_state.jobs
        for j in jobs_view:
            with st.container(border=True):
                st.markdown(f"**[{j['title']}]({j['url']})**")
                left, right = st.columns([3,1])
                with left:
                    st.caption(j.get("source","")) ; st.write(j.get("snippet","")) ; st.caption(f"Why: {j.get('why','')}")
                with right:
                    st.markdown(f"<span class='match-pill'>{j.get('match',0)}% match</span>", unsafe_allow_html=True)
    else:
        st.info("No results yet. Enter a company above and click **Search Jobs**.")

with bulk_tab:
    st.subheader("Bulk Companies ‚Äî scan careers pages")
    st.caption("Paste/upload URLs, or load a preset list from data/*.txt. Use the sidebar Bulk tab for quick scans.")
    txt = st.text_area("Paste URLs (one per line)", height=120)
    up = st.file_uploader("Or upload .txt / .csv (column: url)", type=["txt","csv"])
    preset_labels = [f"data/{os.path.basename(p)}" for p in PRESET_TXT_FILES]
    preset_choice = st.selectbox("Or load a preset list", preset_labels, index=0) if preset_labels else None
    if st.button("Load preset") and preset_choice:
        chosen_fp = PRESET_TXT_FILES[preset_labels.index(preset_choice)]
        try:
            with open(chosen_fp, "r", encoding="utf-8") as f:
                st.session_state.bulk_urls = [l.strip() for l in f if l.strip() and not l.strip().startswith('#')]
            st.success(f"Loaded {len(st.session_state.bulk_urls)} URLs from {os.path.basename(chosen_fp)}")
        except Exception as e:
            st.error(f"Could not load preset: {e}")

    col1, col2 = st.columns(2)
    urls: List[str] = []
    if txt.strip():
        urls += [u.strip() for u in txt.splitlines() if u.strip() and not u.strip().startswith('#')]
    if st.session_state.get("bulk_urls"):
        urls += [u for u in st.session_state.bulk_urls if u]
    if up is not None:
        try:
            if up.name.lower().endswith(".csv"):
                dfu = pd.read_csv(up)
                if "url" in dfu.columns:
                    urls += [str(x) for x in dfu["url"].dropna().tolist()]
                else:
                    urls += [str(x) for x in dfu.iloc[:,0].dropna().tolist()]
            else:
                raw = up.read().decode("utf-8", errors="ignore")
                urls += [u.strip() for u in raw.splitlines() if u.strip() and not u.strip().startswith('#')]
        except Exception as e:
            st.error(f"Could not parse file: {e}")
    if not urls and os.path.exists(BULK_FILE):
        try:
            with open(BULK_FILE, "r", encoding="utf-8") as f:
                urls = [l.strip() for l in f if l.strip() and not l.strip().startswith('#')]
            st.info("Loaded companies from data/companies.txt")
        except Exception:
            pass
    urls = list(dict.fromkeys(urls))

    with col1:
        if st.button("Scan for Careers", type="primary"):
            items = []
            for u in urls[:90]:
                dom = extract_domain(u)
                hits = ddg_find_careers(dom)
                if not hits:
                    items.append({"domain": dom, "careers_url": "(none)", "title": "", "snippet": ""})
                else:
                    for h in hits[:3]:
                        items.append({"domain": dom, "careers_url": h["url"], "title": h["title"], "snippet": h["snippet"]})
            st.session_state.bulk_results = items
            st.success("Bulk results updated.")
    with col2:
        if st.session_state.bulk_results:
            dfb = pd.DataFrame(st.session_state.bulk_results)
            st.download_button("‚¨áÔ∏è Download CSV", data=dfb.to_csv(index=False).encode("utf-8"), file_name="bulk_careers.csv", mime="text/csv")

    if st.session_state.bulk_results:
        dfb = pd.DataFrame(st.session_state.bulk_results)
        st.dataframe(dfb, use_container_width=True, hide_index=True)

with apply_tab:
    st.subheader("Apply ‚Äî Tailor Resume & Open Application")
    if not st.session_state.jobs:
        st.info("Search jobs first, then pick a posting here.")
    else:
        options = [f"{j['title']} ‚Äî {j.get('match',0)}%" for j in st.session_state.jobs]
        idx = st.selectbox("Choose a job", options=range(len(options)), format_func=lambda i: options[i])
        job = st.session_state.jobs[idx]

        # Direct application on employer site
        st.link_button("Open official application", url=job["url"])
        st.markdown("---")

        # Resume upload + preview
        st.markdown("**Upload resume (PDF/DOCX/TXT)**")
        up = st.file_uploader("Resume file", type=["pdf","docx","txt"], key="resume_up")
        if up is not None:
            txt, raw, ext = parse_resume(up)
            st.session_state.resume_text = txt
            st.session_state.resume_bytes = raw
            st.session_state.resume_ext = ext
            with st.expander("Preview extracted text", expanded=True):
                st.text_area("Extracted", value=txt, height=240)

        # Pull job description text (best-effort)
        job_text = fetch_text(job["url"]) or job.get("snippet", "")
        st.text_area("Job description (detected)", value=job_text, height=160)

        use_ai_tailor = st.checkbox("Use AI to tailor (opt-in)", value=False)

        if st.button("Tailor now", type="primary"):
            if not st.session_state.resume_text.strip():
                st.warning("Upload a resume first.")
            else:
                tailored = None
                # --- AI path (optional) ---
                if use_ai_tailor and ai_enabled():
                    try:
                        client = get_ai_client()
                        if client:
                            mdl = (
                                st.secrets.get("openai_model")
                                or os.environ.get("OPENAI_MODEL")
                                or "gpt-4o-mini"
                            )
                            system = (
                                "You are an expert resume editor for early-career candidates. "
                                "Rewrite the resume to target the company and role. "
                                "Preserve truthful content, quantify impact, and align with the job description. "
                                "Output plain text sections."
                            )
                            user = (
                                f"Company: {locals().get('company','')}\n"
                                f"Role: {locals().get('role','')}\n"
                                f"Job description:\n{job_text}\n\n"
                                f"Original resume:\n{st.session_state.resume_text}\n\n"
                                "Return ONLY the revised resume as plain text."
                            )
                            out = client.chat.completions.create(
                                model=mdl,
                                temperature=0.4,
                                messages=[
                                    {"role": "system", "content": system},
                                    {"role": "user", "content": user},
                                ],
                            )
                            tailored = (out.choices[0].message.content or "").strip()
                    except Exception:
                        tailored = None

                # --- Heuristic fallback (keyword highlight) ---
                if not tailored:
                    kws = extract_keywords(job_text, top_k=12)
                    base = st.session_state.resume_text
                    for k in kws:
                        base = re.sub(fr"\b({re.escape(k)})\b", r"**\1**", base, flags=re.I)
                    tailored = base

                st.session_state["tailored"] = tailored
                st.success("Tailored resume ready (preview below).")

        # Preview + save/download
        if (tail := st.session_state.get("tailored")):
            st.text_area("Tailored (preview)", value=tail, height=320)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            fname = f"resume_{(locals().get('company') or 'company').lower().replace(' ','_')}_{ts}.docx"
            full = os.path.join(RESUME_DIR, fname)

            if save_docx_from_text(tail, full):
                append_resume_index({
                    "timestamp": ts,
                    "company": locals().get("company",""),
                    "role": locals().get("role",""),
                    "job_url": job["url"],
                    "file": fname
                })
                with open(full, "rb") as f:
                    st.download_button(
                        "‚¨áÔ∏è Download tailored DOCX",
                        data=f.read(),
                        file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                st.success(f"Saved to databank: data/resumes/{fname}")
            else:
                st.download_button(
                    "‚¨áÔ∏è Download tailored TXT",
                    data=tail.encode("utf-8"),
                    file_name=fname.replace(".docx", ".txt"),
                    mime="text/plain",
                )

with contacts_tab:
    st.subheader("Find Recruiters / HR / University Contacts")
    if st.button("Search Contacts", type="primary"):
        hits = find_contacts(locals().get('company',''), locals().get('role',''))
        if hits:
            for h in hits:
                with st.container(border=True):
                    st.markdown(f"**[{h['title']}]({h['url']})**")
                    st.caption(h.get("source",""))
                    st.write(h.get("snippet",""))
        else:
            st.caption("No contacts found ‚Äî try a different spelling or add 'university recruiting'.")

with qr_tab:
    st.subheader("Create a QR Card for Networking")
    name = st.text_input("Full name")
    title_i = st.text_input("Title (e.g., Student)")
    email = st.text_input("Email")
    phone = st.text_input("Phone")
    org = st.text_input("Organization / School")
    linkedin = st.text_input("LinkedIn URL")
    website = st.text_input("Personal site / portfolio URL")
    mode = st.radio("QR content", ["vCard (contacts app)", "Just my LinkedIn URL"], horizontal=True)
    if st.button("Generate QR", type="primary"):
        if segno is None:
            st.error("QR feature requires 'segno'. Ensure it's in requirements.txt and installed.")
        else:
            buf = io.BytesIO()
            if mode.startswith("vCard"):
                vcf = build_vcard(name, title_i, email, phone, org, linkedin, website)
                segno.make(vcf).save(buf, kind="png", scale=6)
            else:
                url = linkedin or website
                if not url:
                    st.warning("Provide your LinkedIn or a URL.")
                else:
                    segno.make(url).save(buf, kind="png", scale=6)
            if buf.getbuffer().nbytes:
                st.image(buf.getvalue(), caption="Scan me!", use_column_width=False)
                st.download_button("‚¨áÔ∏è Download QR (PNG)", data=buf.getvalue(), file_name="my_qr.png", mime="image/png")

with saved_tab:
    st.subheader("My Saved Items")
    if os.path.exists(INDEX_CSV):
        df = pd.read_csv(INDEX_CSV)
        st.dataframe(df, use_container_width=True, hide_index=True)
        for _, r in df.tail(8).iterrows():
            fp = os.path.join(RESUME_DIR, str(r["file"]))
            if os.path.exists(fp):
                with open(fp, "rb") as f:
                    st.download_button(f"‚¨áÔ∏è {r['file']}", data=f.read(), file_name=str(r["file"]))
    else:
        st.caption("No saved resumes yet.")

with help_tab:
    st.subheader("Setup & Notes")
    st.markdown(
        """
        **Local setup**
        ```bash
        python -m venv .venv && source .venv/bin/activate
        pip install -r requirements.txt
        streamlit run app.py
        ```
        **Streamlit Cloud**: push to GitHub ‚Üí deploy ‚Üí add secrets if using AI.

        **AI toggle**: When off (or no API key), the app never calls OpenAI. You still get search and resume keyword tailoring.

        **Bulk lists**: Put any `.txt` under `data/` (one URL per line). Load them from the sidebar Bulk tab.
        """
    )

# -------------------- Minimal tests (run when Streamlit is missing) --------------------
def _run_tests():
    print("Running basic tests‚Ä¶")
    # canonicalization
    assert canonical_company("jacobs") == "Jacobs Solutions"
    # domain utils
    assert extract_domain("https://www.jacobs.com/careers").startswith("www.jacobs.com")
    assert categorize_domain("https://jobs.lever.co/acme") == "ATS"
    # scoring (should be reasonably high for matching title/keywords/location)
    sample_job = {
        "title": "Civil Engineering Intern",
        "url": "https://jobs.lever.co/acme/123",
        "snippet": "Transportation projects using AutoCAD in Los Angeles",
        "source": "ATS",
    }
    score, why = score_job(sample_job, "Civil Engineering Intern", ["transportation", "AutoCAD"], "Los Angeles")
    assert 50 <= score <= 100, f"unexpected score: {score} ({why})"

    # ---- Shim correctness ----
    cols2 = st.columns(2); assert isinstance(cols2, list) and len(cols2) == 2
    cols3 = st.columns([1, 1, 1]); assert isinstance(cols3, list) and len(cols3) == 3
    tabs3 = st.tabs(["a", "b", "c"]); assert isinstance(tabs3, list) and len(tabs3) == 3
    tabs2 = st.tabs(2); assert isinstance(tabs2, list) and len(tabs2) == 2
    ctx = st.spinner("loading"); assert hasattr(ctx, "__enter__") and hasattr(ctx, "__exit__")

    # ---- Source preference + dedupe behavior ----
    st.session_state["source_pref"] = "Balanced"
    job_ats = {"title": "Software Engineer Intern", "url": "https://ats.example/job1", "snippet": "", "source": "ATS"}
    job_board = {"title": "Software Engineer Intern", "url": "https://board.example/job1", "snippet": "", "source": "Job Board"}
    s_ats, _ = score_job(job_ats, "", [], "")
    s_board, _ = score_job(job_board, "", [], "")
    assert s_ats == s_board, f"Balanced should tie (ATS={s_ats}, Board={s_board})"

    st.session_state["source_pref"] = "Discovery"
    s_ats2, _ = score_job(job_ats, "", [], "")
    s_board2, _ = score_job(job_board, "", [], "")
    assert s_board2 > s_ats2, "Discovery should favor Job Boards"

    kept = dedupe_jobs([job_board, job_ats])
    assert len(kept) == 1 and kept[0]["source"] == "ATS", "Dedupe should prefer ATS for duplicate titles"

    print("All tests passed.")

if __name__ == "__main__":
    if not _STREAMLIT_AVAILABLE or os.environ.get("NON_UI_MODE") == "1":
        _run_tests()
